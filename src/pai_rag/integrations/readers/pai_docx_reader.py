"""Docs parser.

"""
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
from llama_index.core.readers.base import BaseReader
from llama_index.core.schema import Document
from pai_rag.utils.markdown_utils import (
    transform_local_to_oss,
    convert_table_to_markdown,
    is_horizontal_table,
    PaiTable,
)
from docx import Document as DocxDocument
import re
import os
from PIL import Image
import time
from io import BytesIO
from loguru import logger


IMAGE_MAX_PIXELS = 512 * 512


class PaiDocxReader(BaseReader):
    """Read docx files including texts, tables, images.

    Args:
        enable_table_summary (bool):  whether to use table_summary to process tables
    """

    def __init__(
        self,
        enable_table_summary: bool = False,
        oss_cache: Any = None,
    ) -> None:
        self.enable_table_summary = enable_table_summary
        self._oss_cache = oss_cache
        logger.info(
            f"PaiDocxReader created with enable_table_summary : {self.enable_table_summary}"
        )

    def _transform_local_to_oss(
        self, image_blob: bytes, image_filename: str, doc_name: str
    ):
        # 暂时不处理Windows图元文件
        if image_filename.lower().endswith(".emf") or image_filename.lower().endswith(
            ".wmf"
        ):
            return None
        image = Image.open(BytesIO(image_blob))
        return transform_local_to_oss(self._oss_cache, image, doc_name)

    def _convert_paragraph(self, paragraph):
        text = paragraph.text.strip()
        if not text:
            return ""

        # 处理标题
        if paragraph.style.name.startswith("Heading"):
            heading_level = int(
                re.search(r"Heading (\d)", paragraph.style.name).group(1)
            )
            if heading_level > 6:
                heading_level = 6
            return f"{'#' * heading_level} {text}\n\n"

        # 处理普通段落
        return f"{text}\n\n"

    def _get_list_level(self, paragraph):
        indent_levels = {
            "List Paragraph": 0,
            "List Bullet": 1,
            "List Number": 1,
            "List Bullet 2": 2,
            "List Number 2": 2,
            "List Bullet 3": 3,
            "List Number 3": 3,
        }

        # 获取段落的样式名称
        style_name = paragraph.style.name
        # 根据样式名称获取层级
        return indent_levels.get(style_name, 0)

    def _convert_list(self, paragraph, level=0):
        text = paragraph.text.strip()
        if not text:
            return ""

        # 处理无序列表
        if paragraph.style.name.startswith("List"):
            return f"{'-' * level} {text}\n"

        # 处理有序列表
        if paragraph.style.name.startswith("List"):
            return f"{level}. {text}\n"

        return ""

    def _convert_table_to_markdown(self, table, doc_name):
        total_cols = max(len(row.cells) for row in table.rows)

        table_matrix = []
        for row in table.rows:
            table_matrix.append(self._parse_row(row, doc_name, total_cols))
        if is_horizontal_table(table_matrix):
            table = PaiTable(data=table_matrix, row_headers_index=[0])
        else:
            table = PaiTable(data=table_matrix, column_headers_index=[0])
        return convert_table_to_markdown(table, total_cols)

    def _parse_row(self, row, doc_name, total_cols):
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, doc_name).strip()
            row_cells[col_index] = cell_content
            col_index += 1
        return row_cells

    def _parse_cell(self, cell, doc_name):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, doc_name)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(self, paragraph, doc_name):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if not image_id:
                        continue
                    image_part = paragraph.part.rels.get(image_id, None)
                    if image_id:
                        image_blob = image_part.blob
                        image_filename = os.path.basename(image_part.partname)
                        image_url = self._transform_local_to_oss(
                            image_blob, image_filename, doc_name
                        )
                        time_tag = int(time.time())
                        alt_text = f"pai_rag_image_{time_tag}_"
                        image_content = f"![{alt_text}]({image_url})"
                        paragraph_content.append(image_content)

            else:
                paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()

    def convert_document_to_markdown(self, doc_path):
        doc_name = os.path.basename(doc_path).split(".")[0]
        doc_name = doc_name.replace(" ", "_")
        document = DocxDocument(doc_path)
        markdown = []

        paragraphs = document.paragraphs.copy()
        tables = document.tables.copy()

        for element in document.element.body:
            if isinstance(element.tag, str) and element.tag.endswith("p"):  # 段落
                paragraph = paragraphs.pop(0)

                if paragraph.style.name.startswith(
                    "List"
                ) or paragraph.style.name.startswith("List"):
                    current_list_level = self._get_list_level(paragraph)
                    markdown.append(self._convert_list(paragraph, current_list_level))
                else:
                    for run in paragraph.runs:
                        if (
                            hasattr(run.element, "tag")
                            and isinstance(element.tag, str)
                            and run.element.tag.endswith("r")
                        ):
                            drawing_elements = run.element.findall(
                                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                            )
                            for drawing in drawing_elements:
                                blip_elements = drawing.findall(
                                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                                )
                                for blip in blip_elements:
                                    embed_id = blip.get(
                                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                    )
                                    if embed_id:
                                        image_part = document.part.related_parts.get(
                                            embed_id
                                        )
                                        image_blob = image_part.blob
                                        image_filename = os.path.basename(
                                            image_part.partname
                                        )
                                        image_url = self._transform_local_to_oss(
                                            image_blob, image_filename, doc_name
                                        )
                                        time_tag = int(time.time())
                                        alt_text = f"pai_rag_image_{time_tag}_"
                                        image_content = f"![{alt_text}]({image_url})"
                                markdown.append(f"{image_content}\n\n")
                    markdown.append(self._convert_paragraph(paragraph))

            elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # 表格
                table = tables.pop(0)
                markdown.append(self._convert_table_to_markdown(table, None))
                markdown.append("\n\n")

        return "".join(markdown)

    def load_data(
        self,
        file_path: Union[Path, str],
        metadata: bool = True,
        extra_info: Optional[Dict] = None,
    ) -> List[Document]:
        """Loads list of documents from Docx file and also accepts extra information in dict format."""
        return self.load(file_path, metadata=metadata, extra_info=extra_info)

    def load(
        self,
        file_path: Union[Path, str],
        metadata: bool = True,
        extra_info: Optional[Dict] = None,
    ) -> List[Document]:
        """Loads list of documents from Docx file and also accepts extra information in dict format.

        Args:
            file_path (Union[Path, str]): file path of Docx file (accepts string or Path).
            metadata (bool, optional): if metadata to be included or not. Defaults to True.
            extra_info (Optional[Dict], optional): extra information related to each document in dict format. Defaults to None.

        Raises:
            TypeError: if extra_info is not a dictionary.
            TypeError: if file_path is not a string or Path.

        Returns:
            List[Document]: list of documents.
        """

        md_content = self.convert_document_to_markdown(file_path)
        logger.info(f"[PaiDocxReader] successfully processed docx file {file_path}.")
        docs = []
        if metadata:
            if not extra_info:
                extra_info = {}
            doc = Document(text=md_content, extra_info=extra_info)

            docs.append(doc)
        else:
            doc = Document(
                text=md_content,
                extra_info=dict(),
            )
            docs.append(doc)
            logger.info(f"processed doc file {file_path} without metadata")
        logger.info(f"[PaiDocxReader] successfully loaded {len(docs)} nodes.")
        return docs
